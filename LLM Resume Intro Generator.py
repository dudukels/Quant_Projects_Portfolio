#######################################
# LLM Resume Intro Generator
#######################################
'''
Input Resume: "C:\Users\jpkdu\Documents\_Quant Related\quant applications\_automations\John Paul Dulay_resume_Quant_rev9-0.docx"
Input Job Description CSV: "C:\ResumeIntroGeneratorInput.csv"
Output: SQL Server localhost\SQLEXPRESS JPKD..ResumeIntroGeneratorOutputHist

'''
import sys
from openai import OpenAI
from docx import Document
import datetime
from sqlalchemy import create_engine
import pandas as pd


####### access openai API, deploy input prompt and obtain output response from openai model
def LLMResumeIntroGen(JOB_DESCRIPTION =''):

    #### read resume file and convert to text
    RESUME_PATH = r"C:\Users\jpkdu\Documents\_Quant Related\quant applications\_automations\John Paul Dulay_resume_Quant_rev9-0.docx"
    doc = Document(RESUME_PATH)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    resume_text = "\n".join(full_text)

    #### Paste your own previous intro(s) here.
    WRITING_SAMPLE = """
    (1) Capable in modeling & forecasting, time-series, and data & statistical analytics using Python, SQL Server, and Power BI. Have strong analytical and communication skills.
    (2) Passionate in data analytics & predictive modeling for Mortgage Finance industry. Have the quantitative & analytical foundations and transferrable skills for any data science tools/platforms.
    (3) Interested in M&A business analytics. Have transferable experience in Mortgage Finance using Python, SQL, and Power BI. Have strong analytical and communication skills.
    (4) Capable in capital markets research, quantitative modeling & analyses, time-series, and risk & statistical analytics using Python. Have strong analytical and communication skills.
    """

    #### Build your prompt
    prompt = f"""
    Based on the following job description, my current resume, and writing style sample, write a coherent and human 200-250-character resume summary that showcases my relevant skills & experiences and passion for the industry.

    Job Description:
    {JOB_DESCRIPTION}

    My Resume:
    {resume_text}

    Writing Style Sample:
    {WRITING_SAMPLE}

    Always start the response with and fill out the entries in braces: "US Citizen interested in [JobTitle] at [Company]. [Rest of the response]"
    """

    #### Set your OpenAI API key
    api_key = ""
    try: 
        client = OpenAI(api_key=api_key)

        response = client.responses.create(
            model = "gpt-5",
            input = prompt
            #max_tokens=90,
            #temperature=0.8
        )

    except Exception as e:
        print('Error raised in OpenAI script')
        return None

    #### write the LLM response to docx resume
    IntroText = response.output_text 

    ## extract the company name from the response
    # Find index of first period from the left
    left_period_idx = IntroText.find('.')
    # Find the index of the closest space before the period
    previous_space_idx = IntroText.rfind(' ', 0, left_period_idx)
    # Extract the word between the space and the period
    Company = IntroText[previous_space_idx + 1:left_period_idx]

    ## extract the JobTitle from the response
    # Find index of first period from the left
    left_period_idx = IntroText.find(' at ')
    # Find the index of the closest space before the 
    previous_space_idx = IntroText.rfind(' in ', 0, left_period_idx)
    # Extract the word between the space and the
    JobTitle = IntroText[previous_space_idx+3 + 1:left_period_idx]


    #### write the generated intro to a new docx
    OUTPUT_PATH = RESUME_PATH[:-5] + f'_{Company}' + RESUME_PATH[-5:] 

    for para in doc.paragraphs:
        # Detect and replace placeholders
        if "[IntroText]" in para.text:
            para.text = para.text.replace("[IntroText]", IntroText)
    try:
        doc.save(OUTPUT_PATH)
        print(f"Resume saved at the path: {OUTPUT_PATH}")
    except PermissionError:
        print('PermissionError: An instance of the docx is probably opened on MS Word')

    # returns date generated, job title, company, docx path
    return {'Dt' : datetime.date.today(), 
            'JobTitle' : JobTitle,
            'Company' : Company,
            'Path' : OUTPUT_PATH,
            'Intro' : IntroText
           }


def main():
    
    #### read input csv of job description and convert each row of job description to a list
    # all possible encoding for read_csv
    encoding_list = ['ascii','utf_32', 'utf_32_be', 'utf_32_le', 'utf_16', 'utf_16_be', 'utf_16_le', 'utf_7', 'utf_8', 'utf_8_sig', 'latin1', 'cp1252']
    # iterate through each encoder
    for encode in encoding_list:
        try:
            df = pd.read_csv("C:\ResumeIntroGeneratorInput.csv", encoding=encode )

        except Exception as e:
            print(e)
            continue
        else:
            print(f'Success decode using codec {encode}')

    JobDesc_list = df['JobDesc'].tolist()


    #### iterate through each job description and run the python resume job
    output_df = pd.DataFrame() # initialize df
    for JobDesc in JobDesc_list:
        try:
            # generates docx and an output dict
            output_dict = LLMResumeIntroGen(JobDesc)
        except Exception as e:
            print(e)
        else:
            # generate pandas dataframe for recording successful job  
            output_df = pd.concat([output_df, pd.DataFrame( data=output_dict.values(), index=output_dict.keys() ).T ])


    #### append output df to sql server table
    # SQL Server connection info
    server = 'localhost\\SQLEXPRESS' 
    database = 'JPKD'

    # Windows Authentication connection string (no username or password needed)
    connection_string = (
        f"mssql+pyodbc://@{server}/{database}"
        "?driver=ODBC+Driver+17+for+SQL+Server"
        "&trusted_connection=yes"
    )

    try:
        # Create the SQLAlchemy engine
        engine = create_engine(connection_string)

        # Write the DataFrame into SQL Server
        output_df.to_sql(
                'ResumeIntroGeneratorOutputHist',    # Target table name
                con=engine,
                schema='dbo',                        # Use 'dbo' schema
                if_exists='append',                  # 'replace', 'append', or 'fail'
                index=False
        )
    except Exception as e:
        print(e)
    else:
        print('Outputs succesfully recorded in SQL Server: localhost\\SQLEXPRESS JPKD..dbo.ResumeIntroGeneratorOutputHist')

if __name__ == "__main__":
    main()








